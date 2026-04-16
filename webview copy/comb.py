import io
import json
import csv as csv_mod
import xml.etree.ElementTree as ET
from js import document, Blob, URL, Uint8Array
from pyodide.ffi.wrappers import add_event_listener
from pyodide.ffi import create_proxy

# ─── Parser logic (self-contained, ported from scripts/parser) ───

def get_inner_xml(node):
    """XMLノードの内部テキスト（タグを含む）を取得する"""
    if node is None:
        return ""
    return (node.text or "") + "".join(
        ET.tostring(child, encoding='unicode', method='xml') for child in node
    )

def parse_xliff(content_str):
    """XLIFF系ファイル（XLF, MXLIF, MQXLIFF, SDLXLIFF）のパーサー"""
    root = ET.fromstring(content_str)
    items = []
    for tu in root.iter():
        if tu.tag.split('}')[-1] == 'trans-unit':
            source_node = None
            target_node = None
            for child in tu:
                tag_local = child.tag.split('}')[-1]
                if tag_local == 'source':
                    source_node = child
                elif tag_local == 'target':
                    target_node = child
            if source_node is not None:
                src = get_inner_xml(source_node)
                tgt = get_inner_xml(target_node) if target_node is not None else ""
                items.append({'src': src, 'tgt': tgt})
    return items

def parse_tmx(content_str, source_lang='ja', target_lang='en'):
    """TMXファイルのパーサー"""
    root = ET.fromstring(content_str)
    items = []
    for tu in root.findall('.//tu'):
        tuvs = tu.findall('tuv')
        source_seg = ""
        target_seg = ""
        for tuv in tuvs:
            lang = tuv.attrib.get('{http://www.w3.org/XML/1998/namespace}lang')
            if not lang:
                lang = tuv.attrib.get('xml:lang')
            seg = tuv.find('seg')
            if seg is not None:
                text = get_inner_xml(seg)
                if lang and lang.lower().startswith(source_lang.lower()):
                    source_seg = text
                elif lang and lang.lower().startswith(target_lang.lower()):
                    target_seg = text
        if source_seg and target_seg:
            items.append({'src': source_seg, 'tgt': target_seg})
    return items

def parse_docx(content_bytes):
    """DOCXファイルのパーサー（memoQ, Xbench, Phrase 対応）"""
    from docx import Document
    doc = Document(io.BytesIO(content_bytes))
    num_tables = len(doc.tables)
    if num_tables == 0:
        return []

    try:
        if num_tables == 1 and "重要！セグメントIDやソーステキストを変更しないでください" in doc.tables[0].cell(0, 0).text:
            return _extract_memoq(doc)
        if num_tables == 2 and "Exported with ApSIC" in doc.tables[0]._element.xml:
            return _extract_xbench(doc)
        if num_tables > 1 and "When a segment gets repeated" in doc.tables[0].cell(0, 0).text:
            return _extract_phrase(doc)
    except Exception as e:
        raise ValueError(f"DOCX format detection error: {e}")
    
    raise ValueError("Unsupported DOCX format")

def _extract_memoq(doc):
    table = doc.tables[0]
    rows = table.rows
    segments = []
    for i in range(2, len(rows)):
        cells = rows[i].cells
        if len(cells) < 3:
            continue
        src = cells[1].text.strip()
        tgt = cells[2].text.strip()
        note = cells[3].text.strip() if len(cells) > 3 else ""
        if not src and not tgt:
            continue
        seg = {'src': src, 'tgt': tgt}
        if note:
            seg['note'] = note
        segments.append(seg)
    return segments

def _extract_xbench(doc):
    table = doc.tables[1]
    segments = []
    for row in table.rows:
        cells = row.cells
        if len(cells) < 2:
            continue
        src = cells[0].text.strip()
        tgt = cells[1].text.strip()
        if not src and not tgt:
            continue
        segments.append({'src': src, 'tgt': tgt})
    return segments

def _extract_phrase(doc):
    segments = []
    for table in doc.tables[3:]:
        for row in table.rows:
            cells = row.cells
            if len(cells) < 7:
                continue
            src = cells[3].text.strip()
            tgt = cells[4].text.strip()
            note = cells[6].text.strip()
            if not src and not tgt:
                continue
            seg = {'src': src, 'tgt': tgt}
            if note:
                seg['note'] = note
            segments.append(seg)
    return segments

def parse_xlsx(content_bytes):
    """XLSXファイルのパーサー（A=Source, B=Target, C=Note）"""
    import openpyxl
    workbook = openpyxl.load_workbook(io.BytesIO(content_bytes), read_only=True, data_only=True)
    sheet = workbook.active
    if sheet is None:
        return []
    segments = []
    for row in sheet.iter_rows(values_only=True):
        source_val = str(row[0]) if len(row) > 0 and row[0] is not None else ""
        src = source_val.replace('\n', '\\n').strip()
        target_val = str(row[1]) if len(row) > 1 and row[1] is not None else ""
        tgt = target_val.replace('\n', '\\n').strip()
        note_val = str(row[2]) if len(row) > 2 and row[2] is not None else ""
        note = note_val.replace('\n', '\\n').strip()
        if not src and not tgt:
            continue
        seg = {'src': src, 'tgt': tgt}
        if note:
            seg['note'] = note
        segments.append(seg)
    return segments

# ─── Routing ───

TEXT_EXTS = {'.xlf', '.xliff', '.mxliff', '.mqxliff', '.sdlxliff', '.tmx'}
BINARY_EXTS = {'.docx', '.xlsx'}

def route_parse(filename, file_bytes):
    """拡張子に基づいてパーサーをルーテ# ─── Matching Logic ───"""
    # Wait, the previous replacement inserted "# ─── Matching Logic ───" inside route_parse docstring! Let me fix this manually here.
    ext = '.' + filename.rsplit('.', 1)[-1].lower() if '.' in filename else ''
    
    if ext in {'.xlf', '.xliff', '.mxliff', '.mqxliff', '.sdlxliff'}:
        content_str = file_bytes.decode('utf-8')
        return parse_xliff(content_str)
    elif ext == '.tmx':
        content_str = file_bytes.decode('utf-8')
        return parse_tmx(content_str)
    elif ext == '.docx':
        return parse_docx(file_bytes)
    elif ext == '.xlsx':
        return parse_xlsx(file_bytes)
    else:
        raise ValueError(f"Unsupported file extension: {ext}")

# ─── Matching Logic ───

def strip_tags(text):
    import re
    if not isinstance(text, str): return ""
    return re.sub(r'<.*?>|&lt;.*?&gt;', '', text)

def get_tagged_diff(ref_text: str, src_text: str) -> str:
    from difflib import SequenceMatcher
    matcher = SequenceMatcher(None, ref_text, src_text)
    tagged_text = ""
    for tag, i1, i2, j1, j2 in matcher.get_opcodes():
        chunk = src_text[j1:j2]
        if tag == 'equal':
            tagged_text += chunk
        elif tag == 'insert':
            tagged_text += f"[INS]{chunk}[/INS]"
        elif tag == 'replace':
            tagged_text += f"[REPLACE]{chunk}[/REPLACE]"
    return tagged_text

def apply_tm_matching(base_segments, tm_segments):
    for item in tm_segments:
        if 'src_stripped' not in item:
            item['src_stripped'] = strip_tags(item.get('src', ''))

    results = []
    import js
    from pyodide.ffi import to_js

    def process_extract(query, choices, limit=10, score_cutoff=60.0):
        # Convert Python list to JS array proxy, call JS algorithm, and convert back
        choices_js = to_js(choices)
        matches_js = js.processExtractJS(query, choices_js, limit, score_cutoff)
        return matches_js.to_py()
    
    for row in base_segments:
        res_row = dict(row)
        src_orig = row.get('src', '')
        src_stripped = strip_tags(src_orig)
        src_len = len(src_stripped)
        
        candidates = []
        for tm in tm_segments:
            tm_len = len(tm.get('src_stripped', ''))
            if abs(tm_len - src_len) <= (src_len * 0.25):
                candidates.append(tm)
                
        if not candidates or not src_orig:
            results.append(res_row)
            continue
            
        cand_srcs = [c['src_stripped'] for c in candidates]
        matches = process_extract(src_stripped, cand_srcs, limit=10, score_cutoff=60.0)
        
        seen_srcs = set()
        valid_count = 0
        
        for match_text, score, idx in matches:
            tm_item = candidates[idx]
            tm_src = tm_item.get('src', '')
            tm_tgt = tm_item.get('tgt', '')
            
            if tm_src in seen_srcs:
                continue
            seen_srcs.add(tm_src)
            
            if score >= 99.9:
                tagged_src = tm_src
            else:
                tagged_src = get_tagged_diff(tm_src, src_orig)
                
            res_row[f"sim{valid_count+1}_src"] = tagged_src
            res_row[f"sim{valid_count+1}_tgt"] = tm_tgt
            
            valid_count += 1
            if score >= 99.9 or valid_count >= 2:
                break
                
        results.append(res_row)
    return results

# ─── UI Logic & State ───

app_state = {
    "src_files": [],
    "tm_files": []
}
current_segments = []
current_is_matched = False

def chunk_consistency(items, threshold=80.0):
    import js
    
    processed = []
    for item in items:
        new_item = dict(item)
        if 'src_stripped' not in new_item:
            new_item['src_stripped'] = strip_tags(new_item.get('src', ''))
        if 'tgt_stripped' not in new_item:
            new_item['tgt_stripped'] = strip_tags(new_item.get('tgt', ''))
        processed.append(new_item)
        
    results = []
    items_to_process = processed.copy()
    
    while items_to_process:
        seed = items_to_process.pop(0)
        seed_src = seed['src_stripped']
        seed_tgt = seed['tgt_stripped']
        
        temp_chunk = [seed]
        remaining = []
        
        for item in items_to_process:
            if seed_src == item['src_stripped']:
                if seed_tgt == item['tgt_stripped']:
                    continue
                else:
                    temp_chunk.append(item)
            else:
                score = js.fastLevenshteinRatio(seed_src, item['src_stripped'])
                if score >= threshold:
                    temp_chunk.append(item)
                else:
                    remaining.append(item)
        
        # 只取多个或者有差异的 (Optional: uncomment to filter singletons)
        # if len(temp_chunk) > 1:
        results.append(temp_chunk)
        items_to_process = remaining
        
    return results

def set_status(msg, kind="loading"):
    area = document.getElementById("status-area")
    if kind == "loading":
        area.innerHTML = f'<div class="status-msg loading"><span class="spinner"></span>{msg}</div>'
    elif kind == "success":
        area.innerHTML = f'<div class="status-msg success">✓ {msg}</div>'
    elif kind == "error":
        area.innerHTML = f'<div class="status-msg error">✗ {msg}</div>'
    else:
        area.innerHTML = ''

def render_results(segments, is_matched=False):
    global current_segments, current_is_matched
    current_segments = segments
    current_is_matched = is_matched
    
    container = document.getElementById("results-container")
    count_el = document.getElementById("results-count")
    btn_export_csv = document.getElementById("btn-export-csv")
    btn_export_json = document.getElementById("btn-export-json")
    
    if not segments:
        container.innerHTML = '''
            <div class="empty-state">
                <div class="empty-state-icon">📭</div>
                <div class="empty-state-text">セグメントが見つかりませんでした</div>
            </div>
        '''
        count_el.style.display = "none"
        btn_export_csv.disabled = True
        btn_export_json.disabled = True
        return

    has_note = any('note' in s for s in segments)
    
    note_th = '<th>Note</th>' if has_note else ''
    sim_th = '<th>類似文1原文</th><th>類似文1訳文</th><th>類似文2原文</th><th>類似文2訳文</th>' if is_matched else ''
    
    html = f'''
    <div class="results-table-wrap">
        <table>
            <thead><tr><th>#</th><th>Source</th><th>Target</th>{note_th}{sim_th}</tr></thead>
            <tbody>
    '''
    for i, seg in enumerate(segments):
        src_escaped = seg.get('src', '').replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
        tgt_escaped = seg.get('tgt', '').replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
        note_td = ''
        if has_note:
            note_text = seg.get('note', '').replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
            note_td = f'<td class="note-col">{note_text}</td>'
            
        sim_tds = ''
        if is_matched:
            s1_src = seg.get('sim1_src', '').replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
            s1_tgt = seg.get('sim1_tgt', '').replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
            s2_src = seg.get('sim2_src', '').replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
            s2_tgt = seg.get('sim2_tgt', '').replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
            sim_tds = f'<td>{s1_src}</td><td>{s1_tgt}</td><td>{s2_src}</td><td>{s2_tgt}</td>'
            
        html += f'<tr><td class="idx-col">{i+1}</td><td>{src_escaped}</td><td>{tgt_escaped}</td>{note_td}{sim_tds}</tr>'
    
    html += '</tbody></table></div>'
    container.innerHTML = html
    
    count_el.textContent = f"{len(segments)} segments"
    count_el.style.display = "inline"
    btn_export_csv.disabled = False
    btn_export_json.disabled = False

def export_csv(event):
    if not current_segments:
        return
    
    output = io.StringIO()
    writer = csv_mod.writer(output)
    
    has_note = any('note' in s for s in current_segments)
    header = ['#', 'src', 'tgt']
    if has_note:
        header.append('note')
    if current_is_matched:
        header.extend(['類似文1原文', '類似文1訳文', '類似文2原文', '類似文2訳文'])
        
    writer.writerow(header)
    
    for i, seg in enumerate(current_segments):
        row = [i+1, seg.get('src', ''), seg.get('tgt', '')]
        if has_note:
            row.append(seg.get('note', ''))
        if current_is_matched:
            row.extend([
                seg.get('sim1_src', ''), seg.get('sim1_tgt', ''),
                seg.get('sim2_src', ''), seg.get('sim2_tgt', '')
            ])
        writer.writerow(row)
    
    csv_str = output.getvalue()
    blob = Blob.new([csv_str], {"type": "text/csv;charset=utf-8"})
    url = URL.createObjectURL(blob)
    
    a = document.createElement("a")
    a.href = url
    a.download = "parsed_segments.csv"
    a.click()
    URL.revokeObjectURL(url)

def export_json(event):
    if not current_segments:
        return
        
    json_str = json.dumps(current_segments, ensure_ascii=False, indent=2)
    blob = Blob.new([json_str], {"type": "application/json;charset=utf-8"})
    url = URL.createObjectURL(blob)
    
    a = document.createElement("a")
    a.href = url
    a.download = "parsed_segments.json"
    a.click()
    URL.revokeObjectURL(url)

async def read_file(file_obj):
    array_buf = await file_obj.arrayBuffer()
    return route_parse(file_obj.name, array_buf.to_bytes())

# --- Handlers ---

def update_src_display():
    el = document.getElementById("src-file-names")
    if el:
        names = [f.name for f in app_state["src_files"]]
        el.innerText = ", ".join(names) if names else "選択されていません"

def update_tm_display():
    el = document.getElementById("tm-file-names")
    names = [f.name for f in app_state["tm_files"]]
    el.innerText = ", ".join(names) if names else "選択されていません"

async def handle_src_input(e):
    fs = e.target.files
    app_state["src_files"] = [fs.item(i) for i in range(fs.length)]
    update_src_display()

async def handle_src_drop(e):
    e.preventDefault()
    document.getElementById("drop-zone").classList.remove("dragover")
    fs = e.dataTransfer.files
    app_state["src_files"] = [fs.item(i) for i in range(fs.length)]
    update_src_display()

async def handle_tm_input(e):
    fs = e.target.files
    app_state["tm_files"] = [fs.item(i) for i in range(fs.length)]
    update_tm_display()

async def handle_tm_drop(e):
    e.preventDefault()
    document.getElementById("drop-zone-tm").classList.remove("dragover")
    fs = e.dataTransfer.files
    app_state["tm_files"] = [fs.item(i) for i in range(fs.length)]
    update_tm_display()

async def exec_parse_only(e):
    if not app_state["src_files"]:
        set_status("対象ファイルを選択してください", "error")
        return
        
    set_status(f"ファイルを解析中...")
    try:
        segments = []
        for src_f in app_state["src_files"]:
            segments.extend(await read_file(src_f))
        render_results(segments, is_matched=False)
        set_status(f"→ {len(segments)} セグメント抽出完了", "success")
    except Exception as exc:
        set_status(f"エラー: {exc}", "error")

async def exec_parse_match(e):
    if not app_state["src_files"]:
        set_status("対象ファイルを選択してください", "error")
        return
        
    set_status("ファイルを解析・計算中...")
    try:
        base_segments = []
        for src_f in app_state["src_files"]:
            base_segments.extend(await read_file(src_f))
            
        tm_segments = []
        for tm_f in app_state["tm_files"]:
            tm_segments.extend(await read_file(tm_f))
            
        results = apply_tm_matching(base_segments, tm_segments)
        render_results(results, is_matched=True)
        set_status(f"計算完了 → {len(results)} セグメント処理済", "success")
    except Exception as exc:
        set_status(f"エラー: {exc}", "error")

async def exec_consistency(e):
    if not app_state["src_files"]:
        set_status("対象ファイルを選択してください", "error")
        return
        
    set_status(f"ファイルを解析・ゆれ計算中...")
    try:
        segments = []
        for src_f in app_state["src_files"]:
            segments.extend(await read_file(src_f))
            
        chunks = chunk_consistency(segments, threshold=80.0)
        
        flat_results = []
        for i, chunk in enumerate(chunks):
            # ゆれがある（要素が2つ以上）場合のみ表示するなどのフィルタも可能ですが、まずは全て出す
            for item in chunk:
                # 既存のnoteがある場合は追記
                note_orig = item.get('note', '')
                item['note'] = (note_orig + f" [Group {i+1}]").strip()
                flat_results.append(item)
                
        render_results(flat_results, is_matched=False)
        set_status(f"ゆれチェック完了 → {len(chunks)} グループ作成完了", "success")
    except Exception as exc:
        set_status(f"エラー: {exc}", "error")

# ─── Event Binding ───

add_event_listener(document.getElementById("file-input"), "change", handle_src_input)
add_event_listener(document.getElementById("tm-file-input"), "change", handle_tm_input)

add_event_listener(document.getElementById("btn-export-csv"), "click", export_csv)
add_event_listener(document.getElementById("btn-export-json"), "click", export_json)
add_event_listener(document.getElementById("btn-parse-only"), "click", exec_parse_only)
add_event_listener(document.getElementById("btn-parse-match"), "click", exec_parse_match)
add_event_listener(document.getElementById("btn-consistency"), "click", exec_consistency)

dz_src = document.getElementById("drop-zone")
add_event_listener(dz_src, "dragover", lambda e: (e.preventDefault(), dz_src.classList.add("dragover")))
add_event_listener(dz_src, "dragleave", lambda e: dz_src.classList.remove("dragover"))
add_event_listener(dz_src, "drop", handle_src_drop)

dz_tm = document.getElementById("drop-zone-tm")
add_event_listener(dz_tm, "dragover", lambda e: (e.preventDefault(), dz_tm.classList.add("dragover")))
add_event_listener(dz_tm, "dragleave", lambda e: dz_tm.classList.remove("dragover"))
add_event_listener(dz_tm, "drop", handle_tm_drop)

# Hide initial loading overlay
document.getElementById("loading-overlay").classList.add("hidden")
