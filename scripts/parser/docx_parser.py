import io
from typing import List
from docx import Document
from scripts.parser.base import BaseParser
from scripts.util.models import RawDocument
from scripts.util.types import Segment

class DocxParser(BaseParser):
    def parse(self, document: RawDocument) -> List[Segment]:
        """
        Wordドキュメント(docx)をパースしてSegmentのリストを返す。
        """
        if not isinstance(document.contents, bytes):
            # 万が一 str で渡された場合 (通常はないはずだが)
            print(f"Warning: {document.filename} is not in bytes. Trying to encode.")
            content_bytes = document.contents.encode('utf-8')
        else:
            content_bytes = document.contents

        try:
            doc = Document(io.BytesIO(content_bytes))
        except Exception as e:
            print(f"Error opening Word document {document.filename}: {e}")
            return []

        num_tables = len(doc.tables)
        if num_tables == 0:
            return []

        # フォーマットの判定と抽出
        try:
            if num_tables == 1 and "重要！セグメントIDやソーステキストを変更しないでください" in doc.tables[0].cell(0, 0).text:
                return self._extract_memoq(doc)
                
            if num_tables == 2 and "Exported with ApSIC" in doc.tables[0]._element.xml:
                return self._extract_xbench(doc)
                
            if num_tables > 1 and "When a segment gets repeated" in doc.tables[0].cell(0, 0).text:
                return self._extract_phrase(doc)
            
            # Fallback for generic tables
            return self._extract_generic(doc)

        except Exception as e:
            print(f"Error during Word format detection/extraction for {document.filename}: {e}")
            return []

    def _extract_generic(self, doc: Document) -> List[Segment]:
        segments: List[Segment] = []
        for table in doc.tables:
            for row in table.rows:
                row_cells = row.cells
                if len(row_cells) < 2:
                    continue
                
                src = row_cells[0].text.strip().replace('\t', '\\t')
                tgt = row_cells[1].text.strip().replace('\t', '\\t')
                
                if not src and not tgt:
                    continue
                    
                segments.append({'src': src, 'tgt': tgt})
        return segments


    def _extract_memoq(self, doc: Document) -> List[Segment]:
        table = doc.tables[0]
        rows = table.rows
        segments: List[Segment] = []
        
        # 0: Title, 1: Header, 2-: Data
        for i in range(2, len(rows)):
            row_cells = rows[i].cells
            if len(row_cells) < 3:
                continue
                
            # Column 1: Source, Column 2: Target, Column 3: Note
            src = row_cells[1].text.strip().replace('\t', '\\t')
            tgt = row_cells[2].text.strip().replace('\t', '\\t')
            note = row_cells[3].text.strip() if len(row_cells) > 3 else ""
            
            if not src and not tgt:
                continue

            seg: Segment = {'src': src, 'tgt': tgt}
            if note:
                seg['note'] = note
            segments.append(seg)

        return segments

    def _extract_xbench(self, doc: Document) -> List[Segment]:
        table = doc.tables[1]
        segments: List[Segment] = []
        for row in table.rows:
            row_cells = row.cells
            if len(row_cells) < 2:
                continue
            
            src = row_cells[0].text.strip().replace('\t', '\\t')
            tgt = row_cells[1].text.strip().replace('\t', '\\t')
            
            if not src and not tgt:
                continue
                
            segments.append({'src': src, 'tgt': tgt})
        return segments

    def _extract_phrase(self, doc: Document) -> List[Segment]:
        segments: List[Segment] = []
        for table in doc.tables[3:]:
            for row in table.rows:
                row_cells = row.cells
                if len(row_cells) < 7:
                    continue
                    
                src = row_cells[3].text.strip().replace('\t', '\\t')
                tgt = row_cells[4].text.strip().replace('\t', '\\t')
                note = row_cells[6].text.strip()
                
                if not src and not tgt:
                    continue

                seg: Segment = {'src': src, 'tgt': tgt}
                if note:
                    seg['note'] = note
                segments.append(seg)
        return segments
